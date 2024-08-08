from typing import Union
import docx
import xmltodict


class Node:
    def __init__(self, num_prefix: str = '', depth: int = 0, source: Union[str, None] = None,
                 _id = None):
        self._id = _id
        self.parents = {}
        self.source = source
        self.num_prefix = num_prefix
        self.depth = depth
        
        
class ParHandler:
    def __init__(self, par: docx.text.paragraph.Paragraph):
        self.par = par
        self.ctext = par.text.strip()
        self.node = Node()
        self.font_size = self.get_par_font_size(par)
        self.bold = self.get_par_bold_option(par)
        self.xml = xmltodict.parse(par._p.xml, process_namespaces=False)
        self.style_id = par.style.style_id
        self.base_style = par.style.base_style
        self.base_style_id = self.base_style.style_id if self.base_style else None
        self.style_name = par.style.name
        
    def get_par_font_size(self, par:  docx.text.paragraph.Paragraph):
        font_sizes = []
        if par.style.font.size:
            font_sizes.append(par.style.font.size.pt)
        font_sizes += [run.font.size.pt for run in par.runs if run.font.size]
        return max(font_sizes) if font_sizes else None
    
    def get_par_bold_option(self, par:  docx.text.paragraph.Paragraph):
        runs_bold_frac = sum([run.bold is True for run in par.runs]) / (len(par.runs) + 1)
        return par.style.font.bold is True or runs_bold_frac > 0.6
    
    def get_full_text(self):
        if self.node.source not in ('HEADING', 'REGEX', 'APPENDIX') \
            and self.node.num_prefix \
            and 'default' not in self.node.num_prefix:
            return self.node.num_prefix + ' ' + self.ctext
        else:
            return self.ctext


class TableHandler:
    def __init__(self, table: docx.table.Table, src_page_width: int, src_page_height: int,
                 text_cell_min_width: float = 0.8, frame_table_min_hight: float = 0.8,
                 min_frame_columns: int = 7, frame_footer_min_indent: float = 0.82, **kwargs):
        self.table = table
        self.xml = xmltodict.parse(table._element.xml, process_namespaces=False)
        self.height = self.get_table_height(self.xml)
        self.width = self.get_table_width(self.xml)
        # Check page is portrait or album
        self.src_page_width = src_page_width if self.width <= src_page_width else src_page_height
        self.src_page_height = src_page_height if self.width <= src_page_width else src_page_width
        self.cols_count = len(self.table.columns)
        self.rows_count = len(self.table.rows)
        self.text_cell_min_width = text_cell_min_width
        self.frame_table_min_hight = frame_table_min_hight
        self.frame_footer_min_indent = frame_footer_min_indent
        self.min_frame_columns = min_frame_columns
        self.merged = set()
        self.rows = []
        self.text_col_starts = self.cols_count
        self.text_col_ends = -1
        self.text_row_starts = self.rows_count
        self.text_row_ends = -1

        self.investigate()
        self.merge_no_border_cells()
        self.detect_text_cells()
        self.has_frame = self.detect_frame()
        
        # Rifine frame bottom border
        if self.has_frame:
            self.text_row_ends = max(self.text_row_ends, self.get_footer_start_row())
                
    def investigate(self):
        for i, row in enumerate(self.table.rows):
            cells = []
            for j, cell in enumerate(row.cells):
                if cell._element in self.merged:
                    continue

                # Detect merged cells
                rowspan = 1
                colspan = 1
                for next_row in self.table.rows[i + 1:]:
                    if next_row.cells[j]._element == cell._element:
                        rowspan += 1
                    else:
                        break
                for next_cell in row.cells[j + 1:]:
                    if next_cell._element == cell._element:
                        colspan += 1
                    else:
                        break
                if rowspan > 1 or colspan > 1:
                    self.merged.add(cell._element)
                
                cell_handler = CellHandler(
                    cell=cell,
                    rowspan=rowspan,
                    colspan=colspan,
                    x=j,
                    y=i,
                    height=sum(self.rows_heights[i: i + rowspan]),
                    indent_top=sum(self.rows_heights[:i])
                )
                # Detect text (for frames) cell
                cell_handler.is_text = cell_handler.width / self.src_page_width > self.text_cell_min_width
                if cell_handler.is_text:
                    self.text_row_starts = min(self.text_row_starts, i)
                    self.text_row_ends = max(self.text_row_ends, i + rowspan)
                    self.text_col_starts = min(self.text_col_starts, j)
                    self.text_col_ends = max(self.text_col_ends, j + colspan)
            
                # cell_handler.ctext = (
                #         f'DBG [COLS / MIN_COLS {len(self.table.columns)} / {self.min_frame_columns};'
                #         f' PAGE W = {self.src_page_width};'
                #         f' PAGE H = {self.src_page_height};'
                #         f'TABLE W = {self.width}; '
                #         f'TABLE H = {self.height}; '
                #         f'CELL W = {cell_handler.width}; '
                #         f'CELL H = {cell_handler.height}; '
                #         f'INDENT_TOP = {cell_handler.indent_top}; '
                #         f'IS TEXT = {cell_handler.is_text}; '
                #         f'FRAME {(self.text_row_starts, self.text_row_ends, self.text_col_starts, self.text_col_ends)}; '
                #         f'X = {cell_handler.x}: +{cell_handler.colspan}; '
                #         f'Y = {cell_handler.y}: +{cell_handler.rowspan}] '
                #         f'BORDES = {get_cell_nil_borders(cell_handler.xml)} '
                #         + cell_handler.ctext
                #     )
                cells.append(cell_handler)
            self.rows.append(cells)
            

    def merge_no_border_cells(self):
        horizontal_merged_rows = []
        for row in self.rows:
            merged_row = row[:1]
            for cell in row[1:]:
                if 'left' in cell.no_borders and merged_row[-1].rowspan == cell.rowspan:
                    merged_row[-1] = left_join_cells(merged_row[-1], cell)
                else:
                    merged_row.append(cell)
            horizontal_merged_rows.append(merged_row)
            
        vertical_merged_rows = [
            {(cell.x, cell.colspan): cell for cell in row}
            for row in horizontal_merged_rows[:1]
        ]
        for row in horizontal_merged_rows[1:]:
            new_row = {}
            for cell in row:
                try:
                    if 'top' in cell.no_borders and \
                        not vertical_merged_rows[-1][(cell.x, cell.colspan)].ctext:
                        vertical_merged_rows[-1][(cell.x, cell.colspan)] = top_join_cells(
                            vertical_merged_rows[-1][(cell.x, cell.colspan)],
                            cell
                        )
                    else:
                        new_row[(cell.x, cell.colspan)] = cell
                except KeyError:
                    new_row[(cell.x, cell.colspan)] = cell
            if new_row:
                vertical_merged_rows.append(new_row)
        
        self.rows = [
            [cell for _, cell in row.items()]
            for row in vertical_merged_rows
        ]

    def detect_text_cells(self):
        for row in self.rows:
            for cell in row:
                cell.is_text = cell.width / self.src_page_width > self.text_cell_min_width
                if cell.is_text:
                    self.text_row_starts = min(self.text_row_starts, cell.y)
                    self.text_row_ends = max(self.text_row_ends, cell.y + cell.rowspan)
                    self.text_col_starts = min(self.text_col_starts, cell.x)
                    self.text_col_ends = max(self.text_col_ends, cell.x + cell.colspan)
            
    def detect_frame(self):
        # Table hight far from page height and page is portrait
        if (self.height / self.src_page_height) < self.frame_table_min_hight:
            return False
        # Table hcols count lower than min cols count in frame
        if self.cols_count < self.min_frame_columns:
            return False
        return self.text_row_starts >= 0
    
    def get_footer_start_row(self):
        for row in self.rows:
            try:
                if (row[0].indent_top / self.src_page_height) > self.frame_footer_min_indent:
                    return row[0].y
            except IndexError:
                continue
        return 0
                

    def get_table_height(self, xml):
        self.rows_heights = []
        for row in xml['w:tbl']['w:tr']:
            try:
                row_height = int(row['w:trPr']['w:trHeight']['@w:val'])
            except (KeyError, TypeError):
                row_height = 0
            self.rows_heights.append(row_height)
        return sum(self.rows_heights)

    def get_table_width(self, xml):
        try:
            return sum([int(col['@w:w']) for col in xml['w:tbl']['w:tblGrid']['w:gridCol']])
        except (KeyError, TypeError):
            return 0

        
class CellHandler:
    def __init__(self, cell, rowspan: int, colspan: int, x: int, y: int,
                 height: int, indent_top: int):
        self.x = x
        self.y = y
        self.paragraphs = cell.paragraphs
        self.width = 0
        self.xml = xmltodict.parse(cell._element.xml, process_namespaces=False)
        try:
            self.width = int(self.xml ['w:tc']['w:tcPr']['w:tcW']['@w:w'])
        except KeyError:
            self.width = 0
        self.height = height
        self.indent_top = indent_top
        self.is_text = False
        self.rowspan = rowspan
        self.colspan = colspan
        self.no_borders = get_cell_nil_borders(self.xml)
        
    @property
    def ctext(self):
        return '\n'.join([c_par.text.strip() for c_par in self.paragraphs]).strip()
        
        
class TableView:
    def __init__(self, node: Node):
        self.rows = []
        self.node = node
    
    def empty(self):
        n_chars = 0
        for r in self.rows:
            for c in r:
                n_chars += len(c.ctext)
        return n_chars == 0
    
    def clean(self):
        # Reemove empty rows
        self.rows = [
            row for row in self.rows
            if any([len(cell.ctext) > 0 for cell in row])
        ]
        # Reemove empty cols
        left_filled_col = min([cell.x for cell in self.rows[0] if cell.ctext])
        right_filled_col = max([cell.x + cell.colspan for cell in self.rows[0] if cell.ctext])
        self.rows = [
            [cell for cell in row if left_filled_col <= cell.x < right_filled_col]
            for row in self.rows
        ]
    

class DocRoot(ParHandler):
    def __init__(self):
        self.node = Node('[Начало документа]', 1, 'ROOT')
        self.node._id = 'default-start-doc'
        self.ctext = ''


def left_join_cells(cell_1: CellHandler, cell_2: CellHandler):
    cell_1.paragraphs += cell_2.paragraphs
    cell_1.no_borders = cell_1.no_borders.union(cell_2.no_borders)
    cell_1.is_text = any([cell_1.is_text, cell_2.is_text])
    cell_1.colspan += cell_2.colspan
    cell_1.width += cell_2.width
    return cell_1


def top_join_cells(cell_1: CellHandler, cell_2: CellHandler):
    cell_1.paragraphs += cell_2.paragraphs
    cell_1.no_borders = cell_1.no_borders.union(cell_2.no_borders)
    cell_1.is_text = any([cell_1.is_text, cell_2.is_text])
    cell_1.rowspan += cell_2.rowspan
    cell_1.height += cell_2.height
    cell_1.indent_top = max(cell_1.indent_top, cell_2.indent_top)
    return cell_1


def get_cell_nil_borders(c_xml):
    cell_borders = c_xml['w:tc']['w:tcPr'].get('w:tcBorders', {})
    nil_borders = []
    for side in ['top', 'bottom', 'left', 'right']:
        try:
            if cell_borders['w:' + side]['@w:val'] == 'nil':
                nil_borders.append(side)
        except KeyError:
            continue
    return set(nil_borders)
